import os
from pathlib import Path
from docx import Document
from fpdf import FPDF


_UNICODE_MAP = str.maketrans({
    "\u2018": "'", "\u2019": "'",
    "\u201c": '"', "\u201d": '"',
    "\u2013": "-", "\u2014": "--",
    "\u2026": "...",
    "\u00a0": " ",
    "\u00e9": "e", "\u00e8": "e", "\u00ea": "e", "\u00eb": "e",
    "\u00e0": "a", "\u00e2": "a", "\u00e4": "a",
    "\u00ee": "i", "\u00ef": "i",
    "\u00f4": "o", "\u00f6": "o",
    "\u00fb": "u", "\u00fc": "u", "\u00f9": "u",
    "\u00e7": "c",
    "\u00c9": "E", "\u00c0": "A", "\u00c2": "A",
    "\u00b0": " deg",
    "\u00ae": "(R)", "\u00a9": "(C)", "\u2122": "(TM)",
    "\u2022": "-",
    "\u00bd": "1/2", "\u00bc": "1/4", "\u00be": "3/4",
    "\u00d7": "x", "\u00f7": "/",
    "\u2212": "-", "\u00b1": "+/-",
    "\u00a7": "S.",
})


def _sanitize(text: str) -> str:
    text = text.translate(_UNICODE_MAP)
    return text.encode("latin-1", errors="replace").decode("latin-1")


def get_metadata(docx_path: str) -> dict:
    doc = Document(docx_path)
    props = doc.core_properties
    return {
        "file": os.path.basename(docx_path),
        "paragraphs": len(doc.paragraphs),
        "title": props.title or "N/A",
        "author": props.author or "N/A",
        "subject": props.subject or "N/A",
        "created": str(props.created)[:10] if props.created else "N/A",
        "modified": str(props.modified)[:10] if props.modified else "N/A",
        "revision": props.revision or "N/A",
    }


def extract_text(docx_path: str) -> str:
    doc = Document(docx_path)
    lines = [para.text for para in doc.paragraphs]
    return "\n".join(lines)


def _paragraph_style(para) -> str:
    style_name = para.style.name.lower() if para.style else "normal"
    if "heading 1" in style_name:
        return "h1"
    if "heading 2" in style_name:
        return "h2"
    if "heading 3" in style_name or "heading 4" in style_name:
        return "h3"
    return "normal"


def convert_to_pdf(docx_path: str, output_path: str) -> int:
    doc = Document(docx_path)
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.set_margins(20, 20, 20)

    page_width = pdf.w - 40

    style_config = {
        "h1": {"size": 16, "style": "B", "spacing": 6},
        "h2": {"size": 14, "style": "B", "spacing": 5},
        "h3": {"size": 12, "style": "B", "spacing": 4},
        "normal": {"size": 11, "style": "", "spacing": 2},
    }

    paragraph_count = 0
    for para in doc.paragraphs:
        text = _sanitize(para.text.strip())
        style = _paragraph_style(para)
        cfg = style_config[style]

        pdf.set_font("Helvetica", cfg["style"], cfg["size"])

        if text:
            pdf.multi_cell(page_width, cfg["size"] * 0.5 + 2, text)
            paragraph_count += 1
        pdf.ln(cfg["spacing"])

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    pdf.output(output_path)
    return paragraph_count
